# app.py — StudyMind AI (Pastel-Purple Apple-style, Pro Word formatting)
# Date: 2025-11-13 (Asia/Seoul)
# Run: streamlit run app.py
# 🌟 v2.11: 3가지 신규 요청 사항 반영 (v2.8 기반)
# (Word 마인드맵 6.5x9.5 비율 유지 스케일링, Pillow 의존성 추가)

import os
import io
import re
import json
import uuid
import textwrap
from typing import List, Dict, Tuple
from contextlib import contextmanager

import streamlit as st
from PyPDF2 import PdfReader

# ───────── Optional: graphviz for mindmap rendering ─────────
try:
    import graphviz as gv
    GV_AVAILABLE = True
except Exception:
    GV_AVAILABLE = False

# ───────── 🌟 (Req #1 v2.11) Pillow(PIL) 라이브러리 추가 ─────────
try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
    st.error("Pillow 라이브러리가 필요합니다. 'pip install Pillow'를 실행해주세요.")
    st.stop()


# ───────── Word export (python-docx) ─────────
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

TABLE_STYLE_NAME = "Table Grid"

# ───────── Web search & scraping (fallback chain) ─────────
import requests
try:
    from duckduckgo_search import DDGS
    DDG_AVAILABLE = True
except Exception:
    DDG_AVAILABLE = False

try:
    import trafilatura
    TRA_AVAILABLE = True
except Exception:
    TRA_AVAILABLE = False

# ───────── OpenAI SDK (version-safe wrapper) ─────────
from openai import OpenAI
from openai import AuthenticationError, RateLimitError, OpenAIError

# ───────── Ensure run via Streamlit ─────────
try:
    from streamlit.runtime.scriptrunner import get_script_run_ctx
    if get_script_run_ctx() is None:
        print("❗ 이 앱은 'python app.py'가 아니라 'streamlit run app.py' 로 실행해야 합니다.")
        raise SystemExit(1)
except Exception:
    pass

# ───────── App Config ─────────
st.set_page_config(page_title="StudyMind AI", page_icon="🧠", layout="wide")

# ───────── 🌟 Pastel Purple Apple-like CSS (v2.6: 폰트 축소) ─────────
PASTEL_CSS = """
<style>
  :root{
    --accent:#7C3AED; --accent-2:#A78BFA;
    --ink-900:#12121A; --ink-700:#2E2A3B; --ink-500:#6D6A75; --ink-300:#A8A5AE;
    --surface:#ffffff; --surface-2:#F7F6FB; --line:#E9E7F2; --radius:16px;
    --shadow-lg: 0 10px 15px -3px rgba(0, 0, 0, 0.1), 0 4px 6px -2px rgba(0, 0, 0, 0.05);
    --shadow-accent: 0 10px 28px rgba(124,58,237,0.12);
  }
  html, body, [data-testid="stAppViewContainer"], [data-testid="stApp"]{
    font-family:-apple-system,BlinkMacSystemFont,"SF Pro Text","SF Pro Display",
      "Helvetica Neue",Helvetica,Arial,"Segoe UI",Roboto,system-ui,sans-serif !important;
    color:var(--ink-900);
    background: linear-gradient(160deg, #F5F3FF 0%, #FFFFFF 60%);
  }
  h1 {
      color: var(--accent) !important;
      font-weight: 700 !important;
  }
  .block-container{ max-width:1200px !important; padding:2.0rem 1.2rem !important; }
  [data-testid="stSidebar"]{ background:var(--surface); border-right:1px solid var(--line); }
  
  .apple-card{
    background:var(--surface);
    border:1px solid var(--line);
    border-radius:var(--radius);
    padding:1.0rem 1.0rem;
    box-shadow:0 8px 28px rgba(124,58,237,0.08);
    margin-bottom:1.0rem;
    transition: all 0.2s ease-in-out; 
  }
  .apple-card:hover {
      transform: translateY(-3px); 
      box-shadow: var(--shadow-accent); 
  }

  .hero{
    display:flex; align-items:center; gap:.9rem; padding:0.9rem 1.1rem;
    background:linear-gradient(135deg, rgba(167,139,250,.18), rgba(124,58,237,.08));
    border:1px solid var(--line); border-radius:18px;
  }
  .hero-badge{
    background:var(--accent); color:white; font-weight:700; padding:.28rem .6rem; border-radius:999px; font-size:.85rem;
    box-shadow:0 6px 20px rgba(124,58,237,.25);
  }
  .hero-text{ color:var(--ink-700); font-weight:600; letter-spacing:-0.01em;}
  .apple-divider{ height:1px; background:var(--line); margin:1.0rem 0; }
  .stButton > button{
    border-radius:12px; border:1px solid var(--line); background:#fff; color:var(--ink-900);
    padding:0.62rem 0.9rem; transition: all .15s ease; font-weight:600;
  }
  .stButton > button:hover{ box-shadow:0 8px 24px rgba(124,58,237,0.15); transform: translateY(-1px);}
  .stButton > button:focus{ outline:2px solid var(--accent); }
  .btn-primary > button{ background:var(--accent)!important; color:white!important; border-color:transparent!important; }
  textarea, input, .stTextInput input{
    border-radius:12px !important; border:1px solid var(--line) !important; background:var(--surface-2);
  }
  [data-testid="stFileUploader"]{
    border:1px dashed var(--line) !important; border-radius:14px !important; padding:0.9rem !important;
    background:var(--surface-2);
  }
  [data-testid="stDownloadButton"] button{
    border-radius:12px; border:1px solid var(--line); background:#fff; font-weight:600;
  }
  .mcq-container { font-size: 0.85em; } 

  [data-testid="stMarkdownContainer"] h2 {
      font-size: 1.4rem !important;
  }
  [data-testid="stMarkdownContainer"] h3 {
      font-size: 1.0rem !important;
  }
</style>
"""
st.markdown(PASTEL_CSS, unsafe_allow_html=True)

# ───────── 🌟 UI 헬퍼: Context Manager (v2) ─────────
@contextmanager
def apple_card():
    st.markdown('<div class="apple-card">', unsafe_allow_html=True)
    try:
        yield
    finally:
        st.markdown('</div>', unsafe_allow_html=True)

# ───────── Secrets / Keys (원본 유지) ─────────
def _safe_get_secret(name: str) -> str:
    v = os.getenv(name, "")
    if v: return v
    try:
        return st.secrets.get(name, "")
    except Exception:
        return ""

OPENAI_API_KEY = _safe_get_secret("OPENAI_API_KEY")
TAVILY_API_KEY = _safe_get_secret("TAVILY_API_KEY")
if not OPENAI_API_KEY:
    st.error("❗ OPENAI_API_KEY가 없습니다. (.streamlit/secrets.toml 또는 환경변수에 설정 필요)")
    st.stop()


DEFAULT_MODEL = "gpt-4o-mini"
WORDS_PER_PAGE_DEFAULT = 500
client = OpenAI(api_key=OPENAI_API_KEY)

# ───────── OpenAI wrapper (v2.3) ─────────
def chat_complete(model: str, messages: List[Dict], temperature: float = 0.2):
    try:
        return client.chat.completions.create(model=model, messages=messages, temperature=temperature)
    except AttributeError:
        import openai as _oai
        _oai.api_key = OPENAI_API_KEY
        return _oai.ChatCompletion.create(model=model, messages=messages, temperature=temperature)

# ───────── Utils (원본 유지) ─────────
def inject(template: str, content: str) -> str:
    return template.replace("{content}", content or "")

def _strip_fences(txt: str) -> str:
    t = (txt or "").strip()
    t = re.sub(r"^```(?:json)?", "", t).strip()
    t = re.sub(r"```$", "", t).strip()
    return t

def safe_json_loads(t: str):
    t = _strip_fences(t).replace("“", '"').replace("”", '"').replace("’", "'").replace("‘", "'")
    if t and "'" in t and '"' not in t: t = t.replace("'", '"')
    t = re.sub(r",\s*([\}\]])", r"\1", t)
    return json.loads(t)

def read_pdf(file) -> str:
    try:
        reader = PdfReader(file)
        return "\n".join([page.extract_text() or "" for page in reader.pages])
    except Exception:
        return ""

# ───────── Web Search + Extraction (v2.5) ─────────
@st.cache_data(ttl=1800, show_spinner=False)
def tavily_search(query: str, max_results: int = 5) -> List[Dict]:
    if not TAVILY_API_KEY: return []
    try:
        r = requests.post(
            "https://api.tavily.com/search",
            headers={"Content-Type": "application/json"},
            json={"api_key": TAVILY_API_KEY, "query": query, "max_results": max_results, "search_depth": "basic", "include_answer": False},
            timeout=20,
        )
        r.raise_for_status() 
        data = r.json()
        return [{"title": it.get("title",""), "url": it.get("url",""), "snippet": it.get("content","")} for it in data.get("results", []) if it.get("url")]
    except Exception:
        return []

@st.cache_data(ttl=1800, show_spinner=False)
def ddg_search(query: str, max_results: int = 5) -> List[Dict]:
    if not DDG_AVAILABLE: return []
    try:
        with DDGS() as ddgs:
            res = ddgs.text(query, max_results=max_results)
        return [{"title": r.get("title",""), "url": r.get("href","") or r.get("url",""), "snippet": r.get("body","")} for r in (res or []) if (r.get("href") or r.get("url"))]
    except Exception:
        return []

@st.cache_data(ttl=1800, show_spinner=False)
def fetch_page_text(url: str, max_chars: int = 9000) -> str:
    try:
        html = requests.get(url, timeout=15, headers={"User-Agent": "Mozilla/5.0"}).text
        if TRA_AVAILABLE:
            txt = trafilatura.extract(html, include_comments=False, include_tables=False, no_fallback=False)
            if txt and len(txt) > 300: return txt[:max_chars]
        txt = re.sub(r"<[^>]+>", " ", html)
        txt = re.sub(r"\s+", " ", txt)
        return txt[:max_chars]
    except Exception:
        return ""

def summarize_for_queries(text: str, temperature: float = 0.2) -> List[str]:
    prompt = f"""아래 강의/노트 내용을 바탕으로 인터넷 검색용 쿼리 3개를 한국어로 짧게 만들어줘.
조건: 핵심 키워드 중심, 서로 다른 관점/세부 주제, 따옴표·번호 없이 한 줄에 하나.
원문:
{text[:4000]}
"""
    rsp = chat_complete(
        model=DEFAULT_MODEL,
        messages=[{"role":"system","content":"검색어 생성 보조자. 간결·구체·상호다양."},
                  {"role":"user","content":prompt}],
        temperature=temperature,
    )
    out = (rsp.choices[0].message.content or "").strip().splitlines()
    queries = [q.strip("-• ").strip() for q in out if q.strip()]
    return queries[:3] if queries else []

def web_search_and_gather(content: str, pct: int, budget_chars: int = 8000) -> Tuple[str, List[Dict]]:
    if pct <= 0: return "", []
    if pct > 0 and not TAVILY_API_KEY:
        return "", []
        
    queries = summarize_for_queries(content, temperature=0.2)
    picked, used, sources = [], 0, []
    for q in queries:
        hits = tavily_search(q, 5) 
        if not hits and DDG_AVAILABLE:
            hits = ddg_search(q, 5)
            
        for h in hits:
            url = h.get("url", "")
            if not url or any(s["url"] == url for s in sources): continue
            body = fetch_page_text(url)
            if not body or len(body) < 400: continue
            take_budget = max(0, int(budget_chars * (pct / 100.0)) - used)
            if take_budget <= 0: break
            take = body[: min(len(body), take_budget)]
            picked.append(f"[{h.get('title','')}] {url}\n{take}")
            sources.append({"title": h.get("title", ""), "url": url})
            used += len(take)
            if used >= int(budget_chars * (pct / 100.0)): break
        if used >= int(budget_chars * (pct / 100.0)): break
    return "\n\n",sources if not picked else ("\n\n".join(picked), sources)

def gpt_with_web_context(main_prompt: str, content: str, temperature: float) -> str:
    try:
        pct = st.session_state.get("ext_pct", 30)
        web_ctx, _ = web_search_and_gather(content or main_prompt, pct=pct)
        
        merged = f"[인터넷 외부 컨텍스트]\n{web_ctx}\n\n[사용자 제공 자료]\n{inject(main_prompt, content)}" if web_ctx else inject(main_prompt, content)
        
        rsp = chat_complete(
            model=DEFAULT_MODEL,
            messages=[{"role":"system","content":"당신은 신중하고 정확한 학습 비서입니다. 외부 컨텍스트는 참고로만 사용하고, 핵심은 원문에 근거해 구조화합니다."},
                      {"role":"user","content":merged}],
            temperature=temperature,
        )
        return (rsp.choices[0].message.content or "").strip()

    except AuthenticationError:
        st.error("❌ OpenAI API Key 오류: .streamlit/secrets.toml 파일의 OPENAI_API_KEY가 잘못되었거나 유효하지 않습니다.")
        return "" 
    except RateLimitError:
        st.error("❌ OpenAI API 잔액/한도 오류: OpenAI 계정의 크레딧(잔액)이 소진되었거나 월간 사용 한도를 초과했습니다.")
        return "" 
    except OpenAIError as e:
        st.error(f"❌ OpenAI API 일반 오류: {e}")
        return ""
    except Exception as e_general:
        st.error(f"❌ 알 수 없는 생성 오류: {e_general}")
        return ""


# ───────── JSON Fixers (v2.6: 마인드맵 리스트 픽서 추가) ─────────
def fix_cornell_json(text: str) -> dict:
    t = _strip_fences(text).replace("“", '"').replace("”", '"').replace("’", "'").replace("‘", "'")
    if t.lstrip().startswith('"label"') or t.lstrip().startswith("'label'"): t = "{ " + t + " }"
    m = re.search(r"(\{.*\})", t, flags=re.DOTALL)
    if m: t = m.group(1)
    t = re.sub(r",\s*([\}\]])", r"\1", t)
    if "'" in t and '"' not in t: t = t.replace("'", '"')
    obj = json.loads(t)
    obj.setdefault("label", "(파일명 미상)")
    obj.setdefault("title", "강의 요약")
    obj.setdefault("key_terms", [])
    obj.setdefault("notes", [])
    obj.setdefault("summary", "")
    if not isinstance(obj["key_terms"], list): obj["key_terms"] = [str(obj["key_terms"])]
    if not isinstance(obj["notes"], list): obj["notes"] = [str(obj["notes"])]
    return obj

def fix_mindmap_json(text: str) -> dict:
    t = _strip_fences(text).replace("“", '"').replace("”", '"').replace("’", "'").replace("‘", "'")
    if t.lstrip().startswith('"root"') or t.lstrip().startswith("'root'"): t = "{ " + t + " }"
    m = re.search(r"(\{.*\}|\[.*\])", t, flags=re.DOTALL)
    if m: t = m.group(1)
    t = re.sub(r",\s*([\}\]])", r"\1", t)
    if "'" in t and '"' not in t: t = t.replace('"','\\"').replace("'", '"')
    try:
        obj = json.loads(t)
    except Exception:
        m2 = re.search(r"(\{.*\})", t, flags=re.DOTALL)
        obj = json.loads(m2.group(1)) if m2 else {}
    if isinstance(obj, list) and obj and isinstance(obj[0], dict): obj = obj[0]
    if not (isinstance(obj, dict) and "root" in obj): raise ValueError("Mindmap JSON 구조 오류")
    obj.setdefault("children", [])
    return obj

def fix_mindmap_json_list(text: str) -> List[Dict]:
    t = _strip_fences(text).replace("“", '"').replace("”", '"').replace("’", "'").replace("‘", "'")
    t = re.sub(r",\s*([\}\]])", r"\1", t)
    arr = json.loads(t)
    
    if not isinstance(arr, list):
        raise ValueError("AI 응답이 리스트(List) 형식이 아닙니다.")
    
    validated_list = []
    for item in arr:
        if not isinstance(item, dict) or "label" not in item or "map" not in item:
            continue
        try:
            map_str = json.dumps(item["map"])
            map_obj = fix_mindmap_json(map_str)
            item["map"] = map_obj
            validated_list.append(item)
        except Exception:
            continue
    return validated_list


# ───────── 🌟 Word styling helpers (v2: 한글 깨짐 수정) ─────────
def style_document_defaults(doc: Document):
    try:
        style = doc.styles['Normal']
        font = style.font
        font.name = "Gulim"; font.size = Pt(11)
        try:
            style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Gulim')
        except Exception:
            pass 

        pf = style.paragraph_format; pf.space_before = Pt(0); pf.space_after = Pt(6); pf.line_spacing = 1.3
        
        for lvl, sz in [(1, 18), (2, 16), (3, 14)]:
            h = doc.styles[f'Heading {lvl}']
            h.font.name = "Gulim"; h.font.size = Pt(sz)
            try:
                h._element.rPr.rFonts.set(qn('w:eastAsia'), 'Gulim')
            except Exception:
                pass
    except Exception:
        pass 

def ensure_table_style(doc: Document, style_name: str = TABLE_STYLE_NAME):
    try:
        _ = doc.styles[style_name]
    except KeyError:
        pass

def _set_cell_text(cell, text: str, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]; p.alignment = align
    run = p.add_run(str(text) if text is not None else "")
    run.bold = bold; run.font.size = Pt(11); run.font.name = "Gulim"
    try: run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Gulim')
    except Exception: pass

def _style_header_row(row):
    for cell in row.cells:
        props = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), "EEE9FF")
        props.append(shd)
        for p in cell.paragraphs:
            for r in p.runs: r.bold = True

def _set_col_widths(table, widths_inch: List[float]):
    table.autofit = False
    for row in table.rows:
        for i, w in enumerate(widths_inch):
            row.cells[i].width = Inches(w)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

def add_heading_ea(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(level=level)
    run = p.add_run(str(text))
    run.font.name = "Gulim"; run.font.size = Pt(18 if level==1 else 16 if level==2 else 14)
    try: run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Gulim')
    except Exception: pass

# ───────── 🌟 Word exports (v2.11: Pillow 스케일링) ─────────

def _create_base_doc(title_suffix: str) -> Document:
    doc = Document()
    style_document_defaults(doc) 
    add_heading_ea(doc, f"StudyMind AI — {title_suffix}", level=1)
    doc.add_paragraph("")
    ensure_table_style(doc)
    return doc

def _save_doc_to_bytes(doc: Document) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

def make_cornell_docx_per_files(items: List[Tuple[str, dict]]) -> bytes:
    doc = _create_base_doc("코넬식 노트 정리(첨부파일별)")

    for idx, (label, c) in enumerate(items, 1):
        add_heading_ea(doc, f"{idx}. {label}", level=2)
        p = doc.add_paragraph(); r = p.add_run(f"제목: {c.get('title','')}")
        r.font.name="Gulim"; r.bold=True
        doc.add_paragraph("")

        if c.get("key_terms"):
            p = doc.add_paragraph(); r = p.add_run("Key Terms"); r.bold=True
            t = doc.add_table(rows=1, cols=3, style=TABLE_STYLE_NAME)
            _style_header_row(t.rows[0])
            _set_cell_text(t.rows[0].cells[0], "키워드1", True, WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_text(t.rows[0].cells[1], "키워드2", True, WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_text(t.rows[0].cells[2], "키워드3", True, WD_ALIGN_PARAGRAPH.CENTER)
            _set_col_widths(t, [1.2, 2.0, 2.6])
            row=None
            for i, kw in enumerate(c["key_terms"]):
                if i % 3 == 0: row = t.add_row()
                _set_cell_text(row.cells[i % 3], kw)
            doc.add_paragraph("")

        p = doc.add_paragraph(); r = p.add_run("Notes"); r.bold=True
        
        notes_list = c.get("notes", []) or []
        if not notes_list:
            doc.add_paragraph("(내용 없음)")
        
        for note_text in notes_list:
            p_note = doc.add_paragraph(str(note_text), style="List Bullet")
            pf = p_note.paragraph_format
            pf.space_before = Pt(0); pf.space_after = Pt(3) 
            for r in p_note.runs:
                r.font.name = "Gulim"; r.font.size = Pt(11)
                try: r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Gulim')
                except Exception: pass
        
        doc.add_paragraph("") 
        p = doc.add_paragraph(); r = p.add_run("Summary"); r.bold=True
        t = doc.add_table(rows=1, cols=1, style=TABLE_STYLE_NAME)
        _style_header_row(t.rows[0]); _set_cell_text(t.rows[0].cells[0], c.get("summary",""))
        _set_col_widths(t, [5.8])
        doc.add_page_break()

    return _save_doc_to_bytes(doc)

# 🌟 FIX (Req #1, #3 v2.11): 마인드맵 Word 스케일링 (TD, 6.5x9.5 비율유지)
def _mm_to_dot(tree: dict) -> str:
    lines = []
    lines.append('digraph G {')
    # 🌟 (Req #1) rankdir=TD (세로), (Req #3) size="6.5,9.5" (비율 유지하며 박스에 맞춤)
    lines.append('graph [rankdir=TD, fontsize=12, size="6.5,9.5"];')
    lines.append('node [shape=box, style="rounded,filled", fillcolor="#F6F2FF", color="#D6CCFF", fontname="Gulim", fontsize=12];')
    lines.append('edge [color="#C4B5FD", arrowsize=0.7];')
    def walk(node: dict, parent_id: str):
        # 🌟 (Req #3) TD 레이아웃에 맞게 줄바꿈 너비 20으로
        nlabel_raw = node.get("root") or node.get("name") or "노드"
        nlabel = textwrap.fill(nlabel_raw, width=20).replace("\n", "\\n")
        
        nid = "n_" + uuid.uuid4().hex[:8]
        lines.append(f'{nid} [label="{nlabel}"];')
        if parent_id: lines.append(f'{parent_id} -> {nid};')
        for ch in node.get("children", []): walk(ch, nid)
        
    root_label_raw = tree.get("root","주제")
    root_label = textwrap.fill(root_label_raw, width=20).replace("\n", "\\n")
    root_id = "n_" + uuid.uuid4().hex[:8]
    lines.append(f'{root_id} [label="{root_label}"];')
    
    for ch in tree.get("children", []): walk(ch, root_id)
    lines.append("}")
    return "\n".join(lines)

def make_mindmap_docx_per_file(items: List[Dict]) -> bytes:
    doc = _create_base_doc("마인드맵(파일별)")

    for idx, item in enumerate(items, 1):
        label = item.get("label", f"마인드맵 {idx}")
        tree = item.get("map", {"root": "오류", "children": []})
        
        add_heading_ea(doc, f"{idx}. {label}", level=2)
        doc.add_paragraph("")
        
        try:
            if GV_AVAILABLE and PIL_AVAILABLE:
                dot = _mm_to_dot(tree) # 🌟 size="6.5,9.5" 속성이 포함된 dot 생성
                src = gv.Source(dot)
                png_bytes = src.pipe(format="png")
                
                # 🌟 FIX (Req #3 v2.11): Pillow로 이미지 비율 계산
                img_stream = io.BytesIO(png_bytes)
                img = Image.open(img_stream)
                img.close()
                width_px, height_px = img.size
                
                # 0으로 나누기 방지
                if width_px == 0 or height_px == 0:
                    raise ValueError("이미지 크기 0")

                aspect_ratio = float(height_px) / float(width_px)
                
                max_width_in = 6.5
                max_height_in = 9.5
                
                # 너비를 6.5로 고정했을 때의 예상 높이
                target_width = Inches(max_width_in)
                target_height = target_width * aspect_ratio
                
                if target_height > Inches(max_height_in):
                    # 너무 김 -> 높이를 9.5로 고정
                    doc.add_picture(io.BytesIO(png_bytes), height=Inches(max_height_in))
                else:
                    # 적절함 -> 너비를 6.5로 고정
                    doc.add_picture(io.BytesIO(png_bytes), width=Inches(max_width_in))
                
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                doc.add_paragraph("(Graphviz 또는 Pillow 라이브러리가 없어 이미지를 생성할 수 없습니다.)")
        except Exception as e_gv:
            doc.add_paragraph(f"(마인드맵 이미지 생성 실패: {e_gv})")
        
        doc.add_page_break()

    return _save_doc_to_bytes(doc)


def make_quiz_docx(quiz: List[dict], choices: Dict[int, str], score: int) -> bytes:
    doc = _create_base_doc("4지선다 퀴즈 결과")

    for idx, q in enumerate(quiz, 1):
        add_heading_ea(doc, f"문항 {idx}. {q.get('question','')}", level=2)
        t = doc.add_table(rows=1, cols=2, style=TABLE_STYLE_NAME)
        _style_header_row(t.rows[0])
        _set_cell_text(t.rows[0].cells[0], "보기", True, WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(t.rows[0].cells[1], "내용", True)
        _set_col_widths(t, [0.9, 5.1])

        opts = q.get("options", [])
        ans = q.get("answer", "")
        for i, o in enumerate(opts):
            r = t.add_row().cells
            label = chr(ord('A') + i)
            _set_cell_text(r[0], label, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_text(r[1], o, bold=(label == ans))

        sel = choices.get(idx, "")
        doc.add_paragraph(f"선택: {sel if sel else '(미선택)'} / 정답: {ans}")
        exp = q.get("explanation", "")
        if exp: doc.add_paragraph(f"해설: {exp}")
        srcs = q.get("sources") or []
        if srcs:
            doc.add_paragraph("참고 링크:")
            for s in srcs[:5]:
                p = doc.add_paragraph(style=None)
                r = p.add_run(f"- {s.get('title','link')}: {s.get('url','')}")
                r.font.name="Gulim"
                try: r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Gulim')
                except Exception: pass

    doc.add_paragraph("")
    add_heading_ea(doc, f"총점: {score}/{len(quiz)}", level=2)
    
    return _save_doc_to_bytes(doc)

def make_flashcards_docx(cards: List[dict]) -> bytes:
    doc = _create_base_doc("플래시카드")

    for idx, c in enumerate(cards, 1):
        add_heading_ea(doc, f"카드 {idx}: {c.get('front','')}", level=2)
        
        p_ans = doc.add_paragraph()
        p_ans.add_run("정답: ").bold = True
        r_ans = p_ans.add_run(c.get("back",""))
        r_ans.font.name = "Gulim"
        try: r_ans._element.rPr.rFonts.set(qn('w:eastAsia'), 'Gulim')
        except Exception: pass
        
        p_exp = doc.add_paragraph()
        p_exp.add_run("해설: ").bold = True
        for line in (c.get("explain","") or "").splitlines():
            r_exp = p_exp.add_run(line + "\n")
            r_exp.font.name = "Gulim"
            try: r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Gulim')
            except Exception: pass

        srcs = c.get("sources") or []
        if srcs:
            p_src_header = doc.add_paragraph()
            p_src_header.add_run("참고 링크:").bold = True
            for s in srcs[:5]:
                link_text = f"- {s.get('title','link')}: {s.get('url','')}"
                p_link = doc.add_paragraph(link_text, style="List Bullet")
                for r in p_link.runs: 
                    r.font.name = "Gulim"
                    try: r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Gulim')
                    except Exception: pass
        
        p_file = doc.add_paragraph()
        r_file = p_file.add_run(f"(출처 파일: {c.get('from_file','')})")
        r_file.italic = True
        r_file.font.name = "Gulim"
        try: r_file._element.rPr.rFonts.set(qn('w:eastAsia'), 'Gulim')
        except Exception: pass

        if idx < len(cards):
            doc.add_paragraph("---")

    return _save_doc_to_bytes(doc)


def make_markdownish_docx(title: str, text: str, target_pages: int = 3, words_per_page: int = 500) -> bytes:
    doc = _create_base_doc(title) 

    lines = (text or "").splitlines()
    word_count, pages_inserted = 0, 0
    
    in_table = False
    table_data = []

    def flush_table(doc, table_data):
        if not table_data:
            return
        
        cols = len(table_data[0])
        try:
            table = doc.add_table(rows=0, cols=cols, style=TABLE_STYLE_NAME)
            table.autofit = True 
            
            for i, row_cells in enumerate(table_data):
                row = table.add_row()
                is_header = (i == 0) 
                if is_header:
                    _style_header_row(row)
                    
                for j, cell_text in enumerate(row_cells):
                    align = WD_ALIGN_PARAGRAPH.CENTER if is_header else WD_ALIGN_PARAGRAPH.LEFT
                    _set_cell_text(row.cells[j], cell_text, bold=is_header, align=align)
            doc.add_paragraph("") 
        except Exception as e:
            st.warning(f"Word 표 생성 실패: {e}")
        table_data.clear()

    def add_para(content: str, bullet: bool = False):
        nonlocal word_count, pages_inserted
        content = content.rstrip()
        if bullet:
            p = doc.add_paragraph(content, style="List Bullet")
        else:
            p = doc.add_paragraph(content)
        pf = p.paragraph_format
        pf.space_before = Pt(0); pf.space_after = Pt(6); pf.line_spacing = 1.3
        
        for r in p.runs:
            r.font.name = "Gulim"; r.font.size = Pt(11)
            try: r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Gulim')
            except Exception: pass
            
        word_count += len(content.split())
        if word_count >= words_per_page:
            doc.add_page_break()
            word_count = 0; pages_inserted += 1

    for raw in lines:
        s = raw.strip()

        if s.startswith("|") and s.endswith("|"):
            cells = [c.strip() for c in s[1:-1].split("|")]
            if not all(re.match(r"^-+$", c) for c in cells):
                if not in_table:
                    in_table = True 
                    table_data.append(cells)
                elif len(cells) == len(table_data[0]): 
                    table_data.append(cells)
            continue 
        
        if in_table:
            flush_table(doc, table_data)
            in_table = False
            if not s: continue

        if not s:
            doc.add_paragraph("")
            continue
        if s.startswith("### "):
            add_heading_ea(doc, s[4:].strip(), level=3); continue
        if s.startswith("## "):
            add_heading_ea(doc, s[3:].strip(), level=2); continue
        if s.startswith("# "):
            add_heading_ea(doc, s[2:].strip(), level=1); continue
        if re.match(r"^\s*[-*]\s+", s):
            add_para(re.sub(r"^\s*[-*]\s+", "", s), bullet=True)
        else:
            add_para(s)
    
    if in_table:
        flush_table(doc, table_data)

    while pages_inserted < (target_pages - 1):
        doc.add_page_break(); pages_inserted += 1

    return _save_doc_to_bytes(doc)

# ───────── 🌟 Prompts (v2.9: 마인드맵 4단계+) ─────────

PROMPT_CORNELL_JSON = """[매우 중요] 현재보다 3배 더 길고 상세한 내용을 생성해야 합니다.
[매우 중요] **파일별로 각각** 생성해야 합니다. **절대 하나로 합치지 마세요.**
'notes'는 최소 10-15개, 'summary'는 9-15문장으로 구성하세요.

아래 '자료 i: 파일명'별 원문을 각각 코넬식 노트 JSON으로 만들어줘.
반환은 리스트:
[
  {"label":"자료 1: 파일명.pdf","title":"제목",
   "key_terms":["키1","키2","키3", "키4", "키5"],
   "notes":["[필수] 매우 상세한 요점 1 (최소 2-3문장)", "[필수] 매우 상세한 요점 2 (최소 2-3문장)", "... (총 10-15개 이상)"],
   "summary":"[필수] 9~15문장 이상의 매우 상세한 요약"}
  , ...
]
자료:
{content}
"""

# 🌟 FIX (Req #2 v2.9): 마인드맵 3-4단계+ 예시 및 지시어
PROMPT_MINDMAP_JSON_PER_FILE = """[매우 중요] 아래 '자료 i: 파일명'별 원문을 **각각** 별개의 마인드맵 JSON으로 만들어줘.
[매우 중요] **절대 하나로 합치지 마세요.**
[매우 중요] 마인드맵은 **최소 3단계에서 4단계 이상 깊이(root -> 대분류 -> 중분류 -> 소분류...)**로 매우 상세하게 구성하세요.

반환은 반드시 리스트 형식이어야 함:
[
  {"label":"자료 1: 파일명.pdf",
   "map": {
     "root": "자료 1의 핵심 주제",
     "children": [
       {"name":"대분류 1", "children": [
         {"name": "중분류 1-1", "children": [
           {"name": "소분류 1-1-1"},
           {"name": "소분류 1-1-2"}
         ]},
         {"name": "중분류 1-2", "children": [
           {"name": "소분류 1-2-1"}
         ]}
       ]},
       {"name":"대분류 2", "children": [
         {"name": "중분류 2-1"}
       ]}
     ]
   }
  },
  {"label":"자료 2: 노트.txt",
   "map": {
     "root": "자료 2의 핵심 주제",
     "children": [...]
   }
  }
]
자료:
{content}
"""

PROMPT_FLASHCARDS_JSON = """아래 통합 자료(파일명 라벨 포함)를 바탕으로 플래시카드 **정확히 {n}장**을 JSON 배열로 만들어줘.
각 카드에는 다음 필드가 반드시 있어야 함:
- "front": "질문"
- "back": "정답"
- "explain": "**최소 5줄 이상**의 자세한 설명(핵심 개념/근거/비교/예시 포함, 인터넷 자료 반영)"
- "from_file": "어느 자료 출처인지(예: 자료 2: 파일명.pdf)"
- "sources": [중요] **인터넷 검색 결과에 실제 존재하는** 자료만 포함. **절대 가짜 링크 만들지 말 것.** [{"title":"...","url":"..."}]
[
  {"front":"질문","back":"정답","explain":"5줄 이상 상세설명","from_file":"자료 1: ...","sources":[{"title":"...","url":"..."}]}
]
자료:
{content}
"""

PROMPT_QUIZ_JSON = """아래 통합 자료를 바탕으로 4지선다 {n}문항 JSON 생성:
각 문항은 다음 필드를 반드시 포함:
- "question": "질문"
- "options": ["A보기","B보기","C보기","D보기"]
- "answer": "A|B|C|D"
- "explanation": "**최소 5줄 이상**의 매우 자세한 해설(근거·정의·예시·비교) + 가능하면 인터넷 출처 반영"
- "sources": [중요] **인터넷 검색 결과에 실제 존재하는** 자료만 포함. **절대 가짜 링크 만들지 말 것.** [{"title":"...","url":"..."}]
형식:
[
  {"question":"...","options":["...","...","...","..."],"answer":"B","explanation":"5줄 이상","sources":[{"title":"...","url":"..."}]}
]
자료:
{content}
"""

PROMPT_EXAM = """아래 통합 자료를 바탕으로 **현재보다 3배 더 길고 상세한** 시험 대비 핵심 요약(불릿 **20~30개**)과 예상문제 5개를 생성.
각 예상문제에는 **현재보다 3배 더 자세하고 긴 모범답안**(근거/절차/공식/예시 포함, 최소 **12~24문장**)을 함께 제시.
**전체 분량이 현재보다 3배 이상 길어져야 함.**
자료가 파일별로 다른 주제를 포함하면 **[필수] '자료 i: 파일명' 형식의 소제목**을 반드시 달고, 내용은 **통합·대조·정리**하되 출처 구분 명시.
자료:
{content}
"""

PROMPT_EXAM_LONG = """아래 통합 자료를 바탕으로 **[매우 중요] 현재보다 3배 더 길고 매우 상세한 시험 대비 장문 요약**을 작성하세요.
- 첨부 파일이 N개면 각 첨부파일 섹션에 **현재보다 3배 더 많은** 설명 분량 배치(총 최소 **{min_pages} 페이지, {min_words} 단어** 이상).
- **[필수] 반드시 이 분량을 채워야 하며, 각 섹션을 매우 상세하고 길게 설명할 것.**
- 각 첨부파일 섹션은 꼭 **"자료 i: 파일명"** 형식의 제목과 번호를 달 것
- 서로 다른 주제는 파일별로 정리하되 **공통점/차이/상충내용**을 명확히 비교
- 구성(예시):
  1) Executive Summary (핵심 15~25 불릿)
  2) [자료 1: 파일명] **(초 장문)** 상세 요약 (정의/원리/핵심개념/절차/공식/예시/주의점)
  3) [자료 2: 파일명] **(초 장문)** 상세 요약
  4) ...
  5) 상충·혼동 포인트(텍스트 설명 또는 | | | 형식의 표 사용)
  6) 자주 나오는 실수/오개념과 교정 포인트
  7) **예상 문제 5개 + 각 문제의 (현재보다 3배 긴) 자세한 모범답안**(12~24문장, 근거/절차/공식/예시 포함)
  8) 1주/3주/시험 직전 학습 플랜(불릿)
- 한국어, 명확한 소제목/불릿 중심
자료:
{content}
"""

PROMPT_MOOD = """다음 사용자의 오늘 감정/상황을 읽고,
1) 따뜻한 공감 한 문단
2) 원인 가설 2~3개
3) 당장 5분 루틴 (아주 구체적)
4) 내일의 작은 실천 (측정 가능한 행동)
5) 필요시 도움요청 신호와 일반 리소스(한국 기준)
톤: 진심 어린 응원, 과장·설교 금지.
사용자 입력:
{content}
"""

# ───────── Sidebar (v2.5) ─────────
st.sidebar.title("🧠 StudyMind AI")
st.sidebar.caption("Pastel-Purple · Minimal · Focused")

st.sidebar.markdown("### ⚙️ 설정")
ext_pct = st.sidebar.slider("인터넷 외부데이터 활용 퍼센트 %", 0, 100, 30, 10)
st.session_state["ext_pct"] = ext_pct
TEMPERATURE = ext_pct / 100.0
st.sidebar.caption(f"모델: {DEFAULT_MODEL} · temperature={TEMPERATURE:.1f}")

st.sidebar.markdown("---")
page = st.sidebar.radio(
    "📚 메뉴",
    [
        "🔥 시험대비 통합 요약",
        "🌿 마인드맵",
        "💡 플래시카드",
        "🧩 4지선다 퀴즈",
        "📄 코넬식 노트 정리",
        "💚 오늘의 감정 코칭",
    ],
    index=0
)

# ───────── Header + Hero (원본 유지) ─────────
st.header(page) 
st.markdown(
    """
    <div class="hero">
      <span class="hero-badge">STUDYMIND AI</span>
      <span class="hero-text">동서울대학교 <b>스마트드론과</b> · <b>김동엽</b> 제작</span>
    </div>
    """,
    unsafe_allow_html=True
)

st.caption("필요 자료(PDF/TXT 최대 20개) 업로드 또는 노트 붙여넣기 → 아래 기능 카드에서 생성")
st.markdown('<div class="apple-divider"></div>', unsafe_allow_html=True)

# ───────── 🌟 Inputs (v2.9: 동적 스피너용 n_attachments 추가) ─────────
colL, colR = st.columns([1, 1])
with colL:
    with apple_card():
        uploaded_files = st.file_uploader("📎 강의자료 업로드 (PDF/TXT 최대 20개)", type=["pdf","txt"], accept_multiple_files=True)
with colR:
    with apple_card():
        text_input = st.text_area("📝 노션/필기 텍스트",
                                    height=160,
                                    placeholder="수업중 필기한 메모 내용을 붙여주세요")

files = list(uploaded_files) if uploaded_files else []
if len(files) > 20:
    st.warning(f"업로드 파일이 20개를 초과하여 앞의 20개만 사용합니다. (총 {len(files)}개 중 20개 사용)")
    files = files[:20]

attachments: List[Tuple[str, str]] = []
if files:
    for idx, f in enumerate(files, 1):
        name = getattr(f, "name", f"attachment_{idx}")
        body = read_pdf(f) if f.type == "application/pdf" else f.read().decode("utf-8", errors="ignore")
        label = f"자료 {idx}: {name}"
        if body.strip():
            attachments.append((label, body.strip()))
if text_input.strip():
    attachments.append(("사용자 노트", text_input.strip()))

n_attachments = max(1, len(attachments))

parts = [f"### [{label}]\n{txt}" for (label, txt) in attachments]
full_text = "\n\n".join(parts).strip()

if not full_text:
    st.info("자료를 업로드하거나 노트를 붙여넣으면 생성 기능이 활성화됩니다.")

st.markdown('<div class="apple-divider"></div>', unsafe_allow_html=True)


# ───────── Primary button wrapper (원본 유지) ─────────
def primary_button(label: str, key: str = None):
    c = st.container()
    with c:
        st.markdown('<div class="btn-primary">', unsafe_allow_html=True)
        out = st.button(label, key=key, use_container_width=True)
        st.markdown('</div>', unsafe_allow_html=True)
    return out

# ───────── 🌟 Pages (v2.9: 동적 스피너, 입력 없음 경고) ─────────

def show_no_input_warning():
    st.error("⚠️ 자료를 업로드하거나 노트를 붙여넣은 후 버튼을 클릭해주세요.")

# 1) Exam Summary
if page == "🔥 시험대비 통합 요약":
    with apple_card():
        st.caption("• 첨부 + 웹컨텍스트 자동 병합 • 파일별 소제목으로 명확히 구분 • 예상문제 + 자세한 모범답안 포함")
        base_pages = st.slider("📄 목표 페이지 수(가이드)", 3, 10, 3, 1)

        colA, colB = st.columns(2)
        with colA:
            basic_ok = primary_button("통합 요약 생성 (기본)", key="btn_exam_basic")
        with colB:
            long_ok = primary_button("통합 요약 생성 (상세)", key="btn_exam_long")

        if basic_ok:
            if full_text:
                low, high = int(15 * n_attachments), int(30 * n_attachments)
                with st.spinner(f"AI가 기본 요약을 생성 중입니다... (약 {low}~{high}초 소요)"):
                    out = gpt_with_web_context(PROMPT_EXAM, full_text, temperature=TEMPERATURE)
                    if out: 
                        st.session_state["exam_text_basic"] = out
                        st.success("기본 통합 요약 생성 완료!")
                        st.markdown(out)
            else:
                show_no_input_warning() 

        if long_ok:
            if full_text:
                low, high = int(30 * n_attachments), int(60 * n_attachments)
                with st.spinner(f"AI가 상세 요약을 생성 중입니다... (약 {low}~{high}초 소요)"):
                    min_pages = max(base_pages * 3, 1) * 3
                    min_words = min_pages * WORDS_PER_PAGE_DEFAULT
                    long_prompt = PROMPT_EXAM_LONG.format(min_words=min_words, min_pages=min_pages, content="{content}")
                    
                    out = gpt_with_web_context(long_prompt, full_text, temperature=TEMPERATURE)
                    if out: 
                        st.session_state["exam_text_long"] = out
                        st.success("상세 통합 요약 생성 완료!")
                        st.markdown(out)
            else:
                show_no_input_warning() 

        st.markdown('<div class="apple-divider"></div>', unsafe_allow_html=True)
        colX, colY = st.columns(2)
        with colX:
            if st.session_state.get("exam_text_basic"):
                if primary_button("⬇️ (기본) Word 내보내기", key="dl_exam_basic"):
                    docx_bytes = make_markdownish_docx(
                        "시험대비 통합 요약(기본: 예상문제+자세한 답 포함)",
                        st.session_state["exam_text_basic"],
                        target_pages=base_pages,
                        words_per_page=WORDS_PER_PAGE_DEFAULT
                    )
                    st.success("✅ Word 파일 생성 완료! 📄 아래 버튼으로 다운로드하세요.")
                    st.download_button("📥 다운로드", data=docx_bytes,
                        file_name="StudyMind_Exam_Summary_Basic.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True)
        with colY:
            if st.session_state.get("exam_text_long"):
                if primary_button("⬇️ (상세) Word 내보내기", key="dl_exam_long"):
                    docx_bytes = make_markdownish_docx(
                        "시험대비 통합 요약(상세: 예상문제+자세한 답 포함)",
                        st.session_state["exam_text_long"],
                        target_pages=base_pages*3, 
                        words_per_page=WORDS_PER_PAGE_DEFAULT
                    )
                    st.success("✅ Word 파일 생성 완료! 📄 아래 버튼으로 다운로드하세요.")
                    st.download_button("📥 다운로드", data=docx_bytes,
                        file_name="StudyMind_Exam_Summary_Detailed.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True)

# 2) Mindmap
elif page == "🌿 마인드맵":
    with apple_card():
        st.caption("• [신규] 첨부파일마다 각각 마인드맵 생성 • Word 내보내기 지원")
        ok = primary_button("마인드맵 생성 (파일별)", key="btn_mm_per_file")
        
        if ok:
            if full_text:
                low, high = int(15 * n_attachments), int(30 * n_attachments)
                with st.spinner(f"AI가 파일별 마인드맵을 생성 중입니다... (약 {low}~{high}초 소요)"):
                    raw = gpt_with_web_context(PROMPT_MINDMAP_JSON_PER_FILE, full_text, temperature=TEMPERATURE)
                    
                    if raw:
                        try:
                            mm_list = fix_mindmap_json_list(raw)
                            st.session_state["mindmaps_per_file"] = mm_list
                            st.success(f"마인드맵 {len(mm_list)}개 생성 성공!")
                        except Exception as e_json:
                            st.error(f"AI가 응답했으나, JSON 처리 중 오류 발생: {e_json}")
                            st.error(f"AI 원본 응답 (일부): {raw[:500]}...")
            else:
                show_no_input_warning() 

        mm_list = st.session_state.get("mindmaps_per_file", [])
        if mm_list:
            for i, item in enumerate(mm_list, 1):
                label = item.get("label", f"마인드맵 {i}")
                tree = item.get("map", {"root": "오류"})
                
                st.markdown(f"### {i}. {label}")
                st.graphviz_chart(_mm_to_dot(tree)) 
                st.markdown('<div class="apple-divider"></div>', unsafe_allow_html=True)
        
            if primary_button("⬇️ 마인드맵 (파일별) Word 내보내기", key="dl_mm_per_file"):
                docx_bytes = make_mindmap_docx_per_file(mm_list)
                st.success("✅ Word 파일 생성 완료! 📄 아래 버튼으로 다운로드하세요.")
                st.download_button("📥 다운로드", data=docx_bytes,
                                  file_name="StudyMind_Mindmaps_PerFile.docx",
                                  mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                  use_container_width=True)

# 3) Flashcards
elif page == "💡 플래시카드":
    with apple_card():
        st.caption("• 웹컨텍스트 자동 병합 • 질문 클릭 시 답변/해설 펼치기 • Word (텍스트 스타일) 내보내기")
        target_n = st.number_input("생성 개수", 5, 40, 10)
        ok = primary_button("플래시카드 생성", key="btn_cards")
        
        if ok:
            if full_text:
                low, high = int((15 * n_attachments) + target_n), int((30 * n_attachments) + (target_n * 2))
                with st.spinner(f"AI가 플래시카드 {target_n}개를 생성 중입니다... (약 {low}~{high}초 소요)"):
                    debug_context = {"data": ""}
                    
                    try:
                        def gen_cards(n_needed: int) -> List[dict]:
                            n_to_request = n_needed + 1 
                            sub_prompt = PROMPT_FLASHCARDS_JSON.replace("{n}", str(n_to_request))
                            api_result = gpt_with_web_context(sub_prompt, full_text, temperature=TEMPERATURE)
                            if not api_result: 
                                return []
                            debug_context["data"] = api_result 
                            cards = safe_json_loads(api_result)
                            if isinstance(cards, dict): cards = [cards]
                            return cards

                        cards = gen_cards(int(target_n))
                        
                        if cards: 
                            seen = set()
                            unique_cards = []
                            for c in cards:
                                f = (c.get("front","") or "").strip()
                                if f and f not in seen:
                                    unique_cards.append(c)
                                    seen.add(f)
                            cards = unique_cards 

                            if len(cards) < int(target_n):
                                remain = int(target_n) - len(cards)
                                more = gen_cards(remain) 
                                for m in more:
                                    f = (m.get("front","") or "").strip()
                                    if f and f not in seen:
                                        cards.append(m); seen.add(f)
                            
                            cards = cards[:int(target_n)]

                            st.session_state["cards"] = cards
                            st.success(f"플래시카드 {len(cards)}개 생성 완료!")
                        
                    except Exception as e_json:
                        st.error(f"AI가 응답했으나, JSON 처리 중 오류 발생: {e_json}")
                        if debug_context["data"]:
                            st.error(f"AI 원본 응답 (일부): {debug_context['data'][:500]}...")
            else:
                show_no_input_warning() 


        cards = st.session_state.get("cards") or []
        for i, c in enumerate(cards, 1):
            with st.expander(f"카드 {i}: {c.get('front','(질문)')}", expanded=False):
                st.caption(f"출처 파일: {c.get('from_file','(미상)')}")
                st.success(f"정답: {c.get('back','')}")
                if c.get("explain"):
                    st.info(c["explain"])
                srcs = c.get("sources") or []
                if srcs:
                    st.markdown("**참고 링크**")
                    for s in srcs[:5]:
                        title = s.get("title") or s.get("url","link")
                        url = s.get("url","")
                        st.markdown(f"- [{title}]({url})")

        if cards:
            st.markdown('<div class="apple-divider"></div>', unsafe_allow_html=True)
            if primary_button("⬇️ 플래시카드 Word 내보내기", key="dl_cards"):
                docx = make_flashcards_docx(cards)
                st.success("✅ Word 파일 생성 완료! 📄 아래 버튼으로 다운로드하세요.")
                st.download_button("📥 다운로드", data=docx, file_name="StudyMind_Flashcards.docx",
                                  mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                  use_container_width=True)

# 4) MCQ
elif page == "🧩 4지선다 퀴즈":
    with apple_card():
        st.caption("• 옵션 선택 즉시 정·오답및 해설 표시 • Word 내보내기")
        n_q = st.number_input("퀴즈 문항 수", 1, 20, 5, 1)
        ok = primary_button("퀴즈 생성", key="btn_quiz")
        
        if ok:
            if full_text:
                low, high = int((10 * n_attachments) + n_q), int((20 * n_attachments) + (n_q * 2))
                with st.spinner(f"AI가 퀴즈 {n_q}문항을 생성 중입니다... (약 {low}~{high}초 소요)"):
                    prompt = PROMPT_QUIZ_JSON.replace("{n}", str(n_q))
                    data = gpt_with_web_context(prompt, full_text, temperature=TEMPERATURE)
                    
                    if data:
                        try:
                            quiz = safe_json_loads(data)
                            st.session_state["quiz"] = quiz if isinstance(quiz, list) else [quiz]
                            st.session_state["quiz_choices"] = {}
                            st.session_state["quiz_score"] = 0
                            for idx in range(1, len(st.session_state["quiz"]) + 1):
                                st.session_state[f"graded_{idx}"] = False
                                st.session_state[f"correct_{idx}"] = False
                            st.success(f"퀴즈 {len(st.session_state['quiz'])}문항 생성 완료!")
                        except Exception as e_json:
                            st.error(f"AI가 응답했으나, JSON 처리 중 오류 발생: {e_json}")
                            st.error(f"AI 원본 응답 (일부): {data[:500]}...")
            else:
                show_no_input_warning() 

        quiz = st.session_state.get("quiz") or []
        if quiz:
            def _grade_now(idx: int, answer: str, sel: str):
                if not sel: return
                st.session_state["quiz_choices"][idx] = sel
                st.session_state[f"graded_{idx}"] = True
                st.session_state[f"correct_{idx}"] = (sel == answer)

            score = 0
            for idx, q in enumerate(quiz, 1):
                st.markdown(f"### 문항 {idx}. {q.get('question','(질문)')}")
                opts = q.get("options", ["A","B","C","D"])
                answer = q.get("answer", "")
                
                opts_list = [("A",opts[0]),("B",opts[1]),("C",opts[2]),("D",opts[3])]
                
                stored_choice_value = st.session_state.get("quiz_choices", {}).get(idx)
                
                default_index = None 
                if stored_choice_value:
                    for i, (val, label) in enumerate(opts_list):
                        if val == stored_choice_value:
                            default_index = i
                            break
                            
                with st.container():
                    st.markdown('<div class="mcq-container">', unsafe_allow_html=True)
                    sel = st.radio(
                        "보기 선택",
                        options=opts_list,
                        index=default_index, 
                        format_func=lambda x: f"{x[0]}. {x[1]}",
                        key=f"choice_{idx}"
                    )
                    
                    picked = sel[0] if isinstance(sel, tuple) else sel
                    
                    if picked:
                        _grade_now(idx, answer, picked)
                        
                    st.markdown('</div>', unsafe_allow_html=True)

                if st.session_state.get(f"graded_{idx}", False):
                    if st.session_state.get(f"correct_{idx}", False):
                        st.success("정답! ✅"); score += 1
                    else:
                        st.error(f"오답 ❌ (정답: {answer})")
                    exp = q.get("explanation")
                    if exp: st.info(f"해설:\n{exp}")
                    srcs = q.get("sources") or []
                    if srcs:
                        st.markdown("**참고 링크**")
                        for s in srcs[:5]:
                            title = s.get("title") or s.get("url","link")
                            url = s.get("url","")
                            st.markdown(f"- [{title}]({url})")

            st.session_state["quiz_score"] = score
            st.markdown(f"**총점:** {score}/{len(quiz)}")

            st.markdown('<div class="apple-divider"></div>', unsafe_allow_html=True)
            if primary_button("⬇️ 퀴즈 결과 Word 내보내기", key="dl_quiz"):
                docx = make_quiz_docx(quiz, st.session_state.get("quiz_choices", {}), st.session_state.get("quiz_score", 0))
                st.success("✅ Word 파일 생성 완료! 📄 아래 버튼으로 다운로드하세요.")
                st.download_button("📥 다운로드", data=docx, file_name="StudyMind_Quiz_Result.docx",
                                  mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                  use_container_width=True)

# 5) Cornell per-file
elif page == "📄 코넬식 노트 정리":
    with apple_card():
        st.caption("• 첨부파일마다 각각 생성 • Word 표 내보내기")
        ok = primary_button("코넬 노트 생성", key="btn_cornell")
        
        if ok:
            if full_text:
                low, high = int(20 * n_attachments), int(40 * n_attachments)
                with st.spinner(f"AI가 (상세한) 코넬 노트를 생성 중입니다... (약 {low}~{high}초 소요)"):
                    raw = gpt_with_web_context(PROMPT_CORNELL_JSON, full_text, temperature=TEMPERATURE)
                    
                    if raw:
                        try:
                            arr = safe_json_loads(raw)
                            per_items: List[Tuple[str, dict]] = []
                            for obj in arr:
                                fixed = fix_cornell_json(json.dumps(obj))
                                label = fixed.get("label","(파일명 미상)")
                                per_items.append((label, fixed))
                            st.session_state["cornell_per_files"] = per_items
                            st.success(f"코넬 노트 {len(per_items)}개(파일별) 생성 완료!")
                        except Exception as e_json:
                            st.error(f"AI가 응답했으나, JSON 처리 중 오류 발생: {e_json}")
                            st.error(f"AI 원본 응답 (일부): {raw[:500]}...")
            else:
                show_no_input_warning() 


        per_items = st.session_state.get("cornell_per_files") or []
        for i, (label, c) in enumerate(per_items, 1):
            st.markdown(f"### {i}. {label}")
            st.markdown(f"**제목:** {c.get('title','')}")
            st.markdown("**Key Terms:** " + (", ".join(c.get("key_terms",[])) if c.get("key_terms") else "(없음)"))
            st.markdown("**Notes:**")
            for n in c.get("notes", []):
                st.markdown(f"- {n}")
            st.markdown("**Summary:** " + (c.get("summary") or "(요약 없음)"))
            st.markdown('<div class="apple-divider"></div>', unsafe_allow_html=True)

        if per_items:
            if primary_button("⬇️ Word 내보내기(파일별 묶음)", key="dl_cornell"):
                docx_bytes = make_cornell_docx_per_files(per_items)
                st.success("✅ Word 파일 생성 완료! 📄 아래 버튼으로 다운로드하세요.")
                st.download_button("📥 다운로드", data=docx_bytes, file_name="StudyMind_Cornell_PerFiles.docx",
                                  mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                  use_container_width=True)

# 6) Mood Coaching
elif page == "💚 오늘의 감정 코칭":
    with apple_card():
        st.caption("• 공감 → 원인가설 → 5분 루틴 → 내일의 작은 실천 → 리소스")
        mood = st.text_area("오늘의 감정/상황", placeholder="예: 발표 앞두고 불안해요. 어제 잠을 거의 못 잤고…", height=140)
        ok = primary_button("코칭 받기", key="btn_mood")
        
        if ok:
            if mood.strip():
                with st.spinner("AI가 당신의 마음을 읽고 있습니다..."):
                    out = gpt_with_web_context(PROMPT_MOOD, mood.strip(), temperature=TEMPERATURE)
                    if out: 
                        st.markdown(out)
                        st.success("당신은 이미 좋은 방향으로 가고 있어요. 한 걸음씩, 오늘도 충분히 잘하고 있습니다. 🌿")
            else:
                st.error("⚠️ 오늘의 감정/상황을 입력한 후 버튼을 클릭해주세요.")